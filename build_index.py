"""
PhysIQ Knowledge Base Builder — UPGRADED
=========================================
Reads ALL of these from physics_docs/ (including sub-folders):
  • .txt   → plain text notes
  • .pdf   → textbooks, papers, scanned docs (with OCR fallback)
  • .docx  → Word documents
  • .md    → Markdown notes

Usage:
  python3 build_index.py               # normal build
  python3 build_index.py --fast        # skip OCR (faster)
  python3 build_index.py --ocr         # force OCR on every PDF page
  python3 build_index.py --info        # just list files, don't build
"""

import os
import sys
import time
import traceback

# ── CLI flags ──────────────────────────────────────────────────
ARGS        = set(sys.argv[1:])
FAST_MODE   = "--fast" in ARGS
FORCE_OCR   = "--ocr"  in ARGS
INFO_ONLY   = "--info" in ARGS
DOCS_FOLDER = "./physics_docs"
INDEX_FOLDER= "./physics_index"

def banner(msg, char="="):
    print(f"\n{char*55}")
    print(f"  {msg}")
    print(f"{char*55}")

banner("PhysIQ — Knowledge Base Builder  v2.0")
print(f"  Docs folder : {DOCS_FOLDER}")
print(f"  Index folder: {INDEX_FOLDER}")
print(f"  Fast mode   : {FAST_MODE}")
print(f"  Force OCR   : {FORCE_OCR}")


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# 1.  CHECK DOCS FOLDER
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
if not os.path.exists(DOCS_FOLDER):
    print(f"\n❌  '{DOCS_FOLDER}' not found.")
    print("    Create it and put your .txt / .pdf / .docx / .md files inside.")
    sys.exit(1)

# Collect all supported files (recursively)
SUPPORTED = (".txt", ".pdf", ".docx", ".md")
all_files: list[str] = []
for root, dirs, files in os.walk(DOCS_FOLDER):
    for f in sorted(files):
        if f.lower().endswith(SUPPORTED) and not f.startswith("."):
            all_files.append(os.path.join(root, f))

if not all_files:
    print(f"\n❌  No supported files found in '{DOCS_FOLDER}'.")
    print(f"    Add .txt, .pdf, .docx, or .md files and try again.")
    sys.exit(1)

# Group by extension for the summary
by_ext: dict[str, list[str]] = {}
for fp in all_files:
    ext = os.path.splitext(fp)[1].lower()
    by_ext.setdefault(ext, []).append(fp)

print(f"\n📂  Found {len(all_files)} file(s) in '{DOCS_FOLDER}':")
for ext, files in sorted(by_ext.items()):
    print(f"    {ext:6s}  →  {len(files):3d} file(s)")

if INFO_ONLY:
    print("\n  (--info mode — not building index)")
    sys.exit(0)


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# 2.  IMPORT LIBRARIES  (with helpful error messages)
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
print("\n📦  Importing libraries…")
try:
    from langchain.schema import Document
    from langchain.text_splitter import RecursiveCharacterTextSplitter
    from langchain_community.vectorstores import FAISS
    from langchain_community.embeddings import HuggingFaceEmbeddings
    print("    ✅  LangChain / FAISS OK")
except ImportError as e:
    print(f"    ❌  {e}")
    print("    Run:  pip install langchain langchain-community faiss-cpu sentence-transformers")
    sys.exit(1)

try:
    import pypdf
    PYPDF_OK = True
    print("    ✅  pypdf OK")
except ImportError:
    PYPDF_OK = False
    print("    ⚠️   pypdf not found — PDFs will be skipped")
    print("        Run:  pip install pypdf")

DOCX_OK = False
try:
    import docx as _docx
    DOCX_OK = True
    print("    ✅  python-docx OK")
except ImportError:
    print("    ⚠️   python-docx not found — .docx files will be skipped")
    print("        Run:  pip install python-docx")

OCR_OK = False
if not FAST_MODE:
    try:
        import pytesseract
        from PIL import Image
        import fitz  # PyMuPDF
        OCR_OK = True
        print("    ✅  Tesseract / PyMuPDF OK (OCR available for scanned PDFs)")
    except ImportError:
        print("    ℹ️   pytesseract/PyMuPDF not found — scanned PDFs won't be OCR'd")
        print("        Run:  pip install pytesseract pymupdf Pillow")
        print("        (Also install Tesseract: https://github.com/tesseract-ocr/tesseract)")


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# 3.  LOADER FUNCTIONS
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

def load_txt(path: str) -> list[Document]:
    """Load a plain-text or Markdown file."""
    for enc in ("utf-8", "utf-8-sig", "latin-1"):
        try:
            with open(path, encoding=enc) as fh:
                text = fh.read().strip()
            if text:
                return [Document(page_content=text, metadata={"source": path, "type": "text"})]
        except (UnicodeDecodeError, OSError):
            continue
    return []


def load_pdf(path: str) -> list[Document]:
    """
    Load a PDF.
    Strategy:
      1. Try pypdf to extract digital text.
      2. If a page has < 50 chars AND OCR is available, re-process with Tesseract.
      3. Annotate each chunk with page number.
    """
    if not PYPDF_OK:
        return []

    docs_out: list[Document] = []
    try:
        reader = pypdf.PdfReader(path)
        total  = len(reader.pages)
        print(f"        {total} page(s)", end="", flush=True)

        for i, page in enumerate(reader.pages):
            text = (page.extract_text() or "").strip()

            # OCR fallback for image-only pages
            if len(text) < 50 and OCR_OK and not FAST_MODE:
                try:
                    pdf_doc = fitz.open(path)
                    mat     = fitz.Matrix(2, 2)          # 2× zoom for better OCR
                    pix     = pdf_doc[i].get_pixmap(matrix=mat)
                    img     = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                    text    = pytesseract.image_to_string(img).strip()
                    pdf_doc.close()
                    if text:
                        print(" [OCR]", end="", flush=True)
                except Exception:
                    pass

            if text and len(text) > 30:
                docs_out.append(Document(
                    page_content=text,
                    metadata={"source": path, "page": i + 1, "total_pages": total, "type": "pdf"}
                ))

    except Exception as exc:
        print(f"\n    ⚠️   pypdf error on '{path}': {exc}")

    return docs_out


def load_docx(path: str) -> list[Document]:
    """Load a Word document (.docx)."""
    if not DOCX_OK:
        return []
    try:
        import docx as _docx_module
        doc   = _docx_module.Document(path)
        parts = []

        # Paragraphs
        for para in doc.paragraphs:
            t = para.text.strip()
            if t:
                parts.append(t)

        # Tables (extract cell text row by row)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
                if row_text:
                    parts.append(row_text)

        text = "\n".join(parts).strip()
        if text:
            return [Document(page_content=text, metadata={"source": path, "type": "docx"})]
    except Exception as exc:
        print(f"\n    ⚠️   docx error on '{path}': {exc}")
    return []


LOADERS = {
    ".txt":  load_txt,
    ".md":   load_txt,
    ".pdf":  load_pdf,
    ".docx": load_docx,
}


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# 4.  LOAD ALL FILES
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
banner("Step 1 — Loading documents", "-")
all_docs:  list[Document] = []
ok_count   = 0
fail_count = 0

for idx, path in enumerate(all_files, 1):
    fname = os.path.relpath(path, DOCS_FOLDER)
    ext   = os.path.splitext(path)[1].lower()
    loader_fn = LOADERS.get(ext)

    print(f"  [{idx:3d}/{len(all_files)}] {fname} … ", end="", flush=True)
    t0 = time.time()

    if loader_fn is None:
        print("SKIP (unsupported)")
        continue

    try:
        loaded = loader_fn(path)
        if loaded:
            all_docs.extend(loaded)
            elapsed = time.time() - t0
            chars   = sum(len(d.page_content) for d in loaded)
            print(f"✅  {len(loaded)} doc(s), {chars:,} chars  ({elapsed:.1f}s)")
            ok_count += 1
        else:
            print("⚠️   empty / no text extracted")
            fail_count += 1
    except Exception:
        print(f"❌  ERROR")
        traceback.print_exc()
        fail_count += 1

print(f"\n  Loaded : {ok_count} file(s)  ({len(all_docs)} document chunks before splitting)")
print(f"  Skipped: {fail_count} file(s)")

if not all_docs:
    print("\n❌  No content could be loaded.  Aborting.")
    sys.exit(1)


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# 5.  CHUNK
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
banner("Step 2 — Splitting into chunks", "-")
splitter = RecursiveCharacterTextSplitter(
    chunk_size=600,          # slightly larger → more context per chunk
    chunk_overlap=80,
    separators=["\n\n", "\n", ". ", " ", ""],
    length_function=len,
)
chunks = splitter.split_documents(all_docs)
print(f"  ✅  {len(all_docs)} docs  →  {len(chunks)} chunks")
if chunks:
    avg = sum(len(c.page_content) for c in chunks) / len(chunks)
    print(f"  Average chunk size: {avg:.0f} chars")


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# 6.  EMBED
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
banner("Step 3 — Creating embeddings", "-")
print("  Model: sentence-transformers/all-MiniLM-L6-v2")
print("  (Downloads ~90 MB on first run — cached afterwards)")
t0 = time.time()
embeddings = HuggingFaceEmbeddings(
    model_name="all-MiniLM-L6-v2",
    model_kwargs={"device": "cpu"},
    encode_kwargs={"normalize_embeddings": True},
)
print(f"  ✅  Model loaded in {time.time()-t0:.1f}s")


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# 7.  BUILD FAISS INDEX
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
banner("Step 4 — Building FAISS index", "-")
t0 = time.time()
print(f"  Embedding {len(chunks)} chunks … (may take 1-5 minutes)", flush=True)

BATCH = 256
vs = None
for i in range(0, len(chunks), BATCH):
    batch = chunks[i : i + BATCH]
    pct   = min(100, int((i + len(batch)) / len(chunks) * 100))
    bar   = "█" * (pct // 5) + "░" * (20 - pct // 5)
    print(f"  [{bar}] {pct:3d}%  ({i+len(batch)}/{len(chunks)})", end="\r", flush=True)
    if vs is None:
        vs = FAISS.from_documents(batch, embeddings)
    else:
        vs.add_documents(batch)

print(f"\n  ✅  Index built in {time.time()-t0:.1f}s")
vs.save_local(INDEX_FOLDER)
print(f"  ✅  Saved to '{INDEX_FOLDER}/'")


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# 8.  DONE
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
banner("✅  Knowledge base built successfully!")
print(f"  Files processed : {ok_count}")
print(f"  Chunks indexed  : {len(chunks)}")
print(f"  Index saved to  : {INDEX_FOLDER}/")
print()
print("  Next steps:")
print("    streamlit run app_hosted.py")
print()
print("  Tips for adding content:")
print("    • Drop any .pdf textbook into physics_docs/ and re-run")
print("    • Add .docx notes from Word into physics_docs/")
print("    • Create sub-folders inside physics_docs/ — they are scanned too")
print("    • For scanned PDFs, install:  pip install pymupdf pytesseract Pillow")
print("      and also install Tesseract (https://tesseract-ocr.github.io/)")
print()
